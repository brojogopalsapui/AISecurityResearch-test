from __future__ import annotations

from pathlib import Path
from dataclasses import dataclass
from datetime import datetime
from zoneinfo import ZoneInfo
from typing import Tuple
import re
from calendar import month_name
from html import escape

from bs4 import BeautifulSoup
from docx import Document

DATE_STEM_RE = re.compile(r'^(\d{4})-(\d{2})-(.+)$')


@dataclass
class UpdateResult:
    docx_path: Path
    post_id: str
    title: str
    preview: str
    ongoing_path: Path
    index_path: Path
    post_path: Path


def _clean(value) -> str:
    if value is None:
        return ''
    if isinstance(value, str):
        return value.strip()
    return str(value).strip()


def berlin_year_month() -> str:
    return datetime.now(ZoneInfo('Europe/Berlin')).strftime('%Y-%m')


def sanitize_slug(text: str) -> str:
    text = _clean(text).lower()
    text = re.sub(r'[^a-z0-9]+', '-', text)
    text = re.sub(r'-+', '-', text).strip('-')
    return text or 'research-note'




def parse_key_value_paragraphs(doc) -> dict[str, str]:
    metadata: dict[str, str] = {}
    for p in doc.paragraphs[:20]:
        text = _clean(p.text)
        if not text or ':' not in text:
            continue
        key, value = text.split(':', 1)
        key = _clean(key)
        value = _clean(value)
        if key and value:
            metadata[key] = value
    return metadata


def infer_category(docx_path: Path, metadata: dict) -> str:
    raw = _clean(metadata.get('Category'))
    if not raw:
        stem = docx_path.stem.lower()
        if stem.startswith('academic-'):
            raw = 'Academic'
        elif stem.startswith('industry-'):
            raw = 'Industry'
        elif stem.startswith('company-release-'):
            raw = 'Company & Release'

    normalized = raw.lower()
    if 'academic' in normalized:
        return 'Academic'
    if 'industry' in normalized:
        return 'Industry'
    if 'company' in normalized or 'release' in normalized or 'ecosystem' in normalized:
        return 'Company & Release'
    return raw or 'Academic'


def category_to_stream_anchor(category: str) -> str:
    normalized = _clean(category).lower()
    if 'industry' in normalized:
        return 'industry-innovation'
    if 'company' in normalized or 'release' in normalized or 'ecosystem' in normalized:
        return 'companies-releases'
    return 'academic-signals'

def ensure_dated_stem(stem: str, fallback_title: str = '') -> str:
    stem = _clean(stem)
    if DATE_STEM_RE.match(stem):
        return stem
    slug = sanitize_slug(stem or fallback_title)
    return f"{berlin_year_month()}-{slug}"


def stem_to_month_year(stem: str) -> str:
    m = re.match(r'^(\d{4})-(\d{2})-', stem)
    if not m:
        return stem
    year = int(m.group(1))
    month = int(m.group(2))
    if 1 <= month <= 12:
        return f"{month_name[month]} {year}"
    return stem


def detect_doctype(html_text: str) -> str:
    m = re.match(r'\s*(<!DOCTYPE[^>]+>)', html_text, flags=re.IGNORECASE)
    return m.group(1) if m else '<!DOCTYPE html>'


def parse_weekly_docx(path: Path) -> dict:
    doc = Document(str(path))
    data = {'metadata': parse_key_value_paragraphs(doc)}

    if doc.tables:
        table = doc.tables[0]
        for row in table.rows[1:]:
            if len(row.cells) < 2:
                continue
            key = _clean(row.cells[0].text)
            val = _clean(row.cells[1].text)
            if key:
                data['metadata'][key] = val

    wanted_headings = [
        'Preview',
        'Full Note Paragraph 1',
        'Full Note Paragraph 2',
        'What Is Changing Technically',
        'What Reviewers Should Notice',
        'Current Research Tension',
    ]

    paragraphs = doc.paragraphs
    for i, p in enumerate(paragraphs):
        heading = _clean(p.text)
        if heading in wanted_headings:
            j = i + 1
            content = []
            while j < len(paragraphs):
                t = _clean(paragraphs[j].text)
                if t in wanted_headings or t in {'Metadata', 'How the notebook uses this DOCX'}:
                    break
                if t:
                    content.append(t)
                j += 1

            if heading in {'What Is Changing Technically', 'What Reviewers Should Notice'}:
                data[heading] = content
            else:
                data[heading] = '\n'.join(content)

    return data


def pick_latest_docx(folder: Path, template_hints: set[str]) -> Path:
    candidates = []
    for f in folder.glob('*.docx'):
        if f.name in template_hints:
            continue
        candidates.append(f)
    if not candidates:
        raise FileNotFoundError(f'No DOCX files found in {folder}')
    return max(candidates, key=lambda f: f.stat().st_mtime)


def normalize_docx_data(data: dict, docx_path: Path, force_filename_post_id: bool = True) -> Tuple[list[str], str]:
    md = data['metadata']
    warnings: list[str] = []
    title = _clean(md.get('Title')) or _clean(data.get('Preview')) or docx_path.stem
    normalized_stem = ensure_dated_stem(docx_path.stem, title)

    if normalized_stem != docx_path.stem:
        warnings.append(
            f"DOCX filename '{docx_path.name}' is not dated. Generated files will use '{normalized_stem}'."
        )

    if force_filename_post_id:
        existing_post_id = _clean(md.get('Post ID'))
        if existing_post_id and existing_post_id != normalized_stem:
            warnings.append(
                f"Overriding DOCX Post ID '{existing_post_id}' with filename-based Post ID '{normalized_stem}'."
            )
        md['Post ID'] = normalized_stem
    elif not _clean(md.get('Post ID')):
        md['Post ID'] = normalized_stem

    expected_post_link = f'posts/{normalized_stem}.html'
    existing_link = _clean(md.get('Full Post Link (optional)'))
    if existing_link and existing_link != expected_post_link:
        warnings.append(
            f"Overriding DOCX Full Post Link '{existing_link}' with filename-based link '{expected_post_link}'."
        )
    md['Full Post Link (optional)'] = expected_post_link

    if not _clean(md.get('Title')):
        md['Title'] = docx_path.stem.replace('-', ' ').title()
        warnings.append('Title was blank in the DOCX, so one was inferred from the filename.')

    if not _clean(md.get('Meta Line')):
        md['Meta Line'] = f"Research Watch • {md['Title']} • 2026 signal"
        warnings.append('Meta Line was blank in the DOCX, so one was inferred.')

    md['Category'] = infer_category(docx_path, md)
    stream_anchor = category_to_stream_anchor(md['Category'])
    if not _clean(md.get('Related Static Page (optional)')):
        md['Related Static Page (optional)'] = f'../ongoing-work.html#{stream_anchor}'
    if not _clean(md.get('Related Static Page Label (optional)')):
        if stream_anchor == 'academic-signals':
            md['Related Static Page Label (optional)'] = 'Academic stream'
        elif stream_anchor == 'industry-innovation':
            md['Related Static Page Label (optional)'] = 'Industry stream'
        else:
            md['Related Static Page Label (optional)'] = 'Companies & Releases stream'

    if not _clean(data.get('Preview')) and _clean(md.get('Preview')):
        data['Preview'] = _clean(md.get('Preview'))

    return warnings, normalized_stem


def add_link_html(links: list[str], href: str, label: str) -> None:
    href = _clean(href)
    label = _clean(label)
    if not href or not label:
        return
    if href.startswith('http://') or href.startswith('https://'):
        links.append(
            f'<a href="{escape(href, quote=True)}" rel="noopener noreferrer" target="_blank">{escape(label)}</a>'
        )
    else:
        links.append(f'<a href="{escape(href, quote=True)}">{escape(label)}</a>')


def build_watch_article(data: dict, docx_stem: str):
    md = data['metadata']
    post_id = _clean(md.get('Post ID')) or docx_stem
    meta_line = _clean(md.get('Meta Line')) or f'Research Watch • {stem_to_month_year(docx_stem)}'
    title = _clean(md.get('Title')) or 'Untitled research note'
    preview = _clean(data.get('Preview')) or _clean(md.get('Preview'))
    p1 = _clean(data.get('Full Note Paragraph 1'))
    p2 = _clean(data.get('Full Note Paragraph 2'))
    tech_list = data.get('What Is Changing Technically', [])
    reviewer_list = data.get('What Reviewers Should Notice', [])
    tension = _clean(data.get('Current Research Tension'))

    links: list[str] = []
    add_link_html(links, md.get('Full Post Link (optional)'), 'Read full post')
    add_link_html(links, md.get('Related Static Page (optional)'), md.get('Related Static Page Label (optional)'))
    add_link_html(links, md.get('External Link 1 URL (optional)'), md.get('External Link 1 Label (optional)'))
    add_link_html(links, md.get('External Link 2 URL (optional)'), md.get('External Link 2 Label (optional)'))

    tech_items = '\n'.join(
        [f'                      <li>{escape(_clean(item))}</li>' for item in tech_list if _clean(item)]
    )
    reviewer_items = '\n'.join(
        [f'                      <li>{escape(_clean(item))}</li>' for item in reviewer_list if _clean(item)]
    )
    links_html = ' '.join(links)

    article_html = f'''
<article class="watch-note accordion" id="{escape(post_id, quote=True)}">
  <button aria-expanded="false" class="accordion-trigger" type="button">
    <span class="accordion-meta">{escape(meta_line)}</span>
    <span class="accordion-title">{escape(title)}</span>
    <span class="accordion-preview">
      {escape(preview)}
    </span>
    <span class="accordion-cta">Read full note</span>
    <span aria-hidden="true" class="accordion-icon"></span>
  </button>

  <div aria-hidden="true" class="accordion-panel">
    <div class="accordion-panel-inner">
      <p>{escape(p1)}</p>
      <p>{escape(p2)}</p>

      <div class="watch-columns">
        <div class="watch-block">
          <h4>What is changing technically</h4>
          <ul>
{tech_items}
          </ul>
        </div>

        <div class="watch-block">
          <h4>What reviewers should notice</h4>
          <ul>
{reviewer_items}
          </ul>
        </div>
      </div>

      <div class="watch-bottom-note">
        <strong>Current research tension:</strong> {escape(tension)}
      </div>

      <div class="watch-inline-links">
        {links_html}
      </div>
    </div>
  </div>
</article>
'''.strip()
    return post_id, title, preview, article_html


def build_full_post_html(data: dict, docx_stem: str) -> str:
    md = data['metadata']
    post_id = _clean(md.get('Post ID')) or docx_stem
    title = _clean(md.get('Title')) or 'Untitled research note'
    preview = _clean(data.get('Preview')) or _clean(md.get('Preview'))
    p1 = _clean(data.get('Full Note Paragraph 1'))
    p2 = _clean(data.get('Full Note Paragraph 2'))
    tech_list = data.get('What Is Changing Technically', [])
    reviewer_list = data.get('What Reviewers Should Notice', [])
    tension = _clean(data.get('Current Research Tension'))
    month_year = stem_to_month_year(docx_stem)

    related_links: list[str] = []
    add_link_html(related_links, md.get('Related Static Page (optional)'), md.get('Related Static Page Label (optional)'))
    add_link_html(related_links, md.get('External Link 1 URL (optional)'), md.get('External Link 1 Label (optional)'))
    add_link_html(related_links, md.get('External Link 2 URL (optional)'), md.get('External Link 2 Label (optional)'))
    related_links_html = ' '.join(related_links) if related_links else ''
    related_links_block = (
        '<div class="btn-row" style="margin-top:1rem;">' + related_links_html + '</div>' if related_links_html else ''
    )

    tech_items = '\n'.join([f'              <li>{escape(_clean(item))}</li>' for item in tech_list if _clean(item)])
    reviewer_items = '\n'.join([f'              <li>{escape(_clean(item))}</li>' for item in reviewer_list if _clean(item)])

    return f'''<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
  <title>{escape(title)} | Brojogopal Sapui</title>
  <meta name="description" content="{escape(preview)}" />
  <link rel="preconnect" href="https://fonts.googleapis.com" />
  <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin />
  <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap" rel="stylesheet" />
  <link rel="stylesheet" href="../assets/css/style.css" />
</head>
<body>
  <header class="site-header">
    <div class="container nav-wrap">
      <a class="brand" href="../index.html" aria-label="Brojogopal Sapui Home">B<span>S</span></a>
      <nav class="nav">
        <a href="../index.html">Home</a>
        <a href="../about.html">About</a>
        <a href="../research.html">Research</a>
        <a href="../ongoing-work.html">Trending Topics</a>
        <a href="../publications.html">Resources</a>
        <a href="../contact.html">Contact</a>
      </nav>
      <button class="menu-btn" aria-label="Toggle menu" aria-expanded="false">
        <span></span>
        <span></span>
        <span></span>
      </button>
    </div>
  </header>

  <main>
    <section class="page-hero">
      <div class="container">
        <span class="eyebrow">Research Watch • {escape(month_year)}</span>
        <h1>{escape(title)}</h1>
        <p class="lead">{escape(preview)}</p>
      </div>
    </section>

    <section class="section">
      <div class="container split">
        <div class="content-card">
          <span class="kicker">Overview</span>
          <h2>What is changing</h2>
          <p>{escape(p1)}</p>
          <p>{escape(p2)}</p>
        </div>

        <div class="content-card">
          <span class="kicker">Why it matters</span>
          <h2>Research significance</h2>
          <ul class="check-list">
{tech_items}
          </ul>
        </div>
      </div>
    </section>

    <section class="section alt">
      <div class="container">
        <div class="section-head">
          <span class="eyebrow">Discussion</span>
          <h2>What reviewers should notice</h2>
          <p>
            These review points help separate benchmark-level claims from stronger system-level conclusions.
          </p>
        </div>

        <div class="content-card">
          <ul class="check-list">
{reviewer_items}
          </ul>
        </div>
      </div>
    </section>

    <section class="section">
      <div class="container">
        <div class="content-card">
          <span class="kicker">Current research tension</span>
          <h2>Why this topic matters now</h2>
          <p>{escape(tension)}</p>
          {related_links_block}
        </div>
      </div>
    </section>

    <section class="section alt">
      <div class="container">
        <div class="cta">
          <div>
            <span class="eyebrow">Next Step</span>
            <h2>Back to ongoing research updates</h2>
            <p>
              Return to the running list of research-watch topics and evolving system-level notes.
            </p>
          </div>
          <div class="cta-actions">
            <a class="btn btn-primary" href="../ongoing-work.html#{escape(post_id, quote=True)}">Back to this note</a>
            <a class="btn btn-secondary" href="../research.html">Research</a>
          </div>
        </div>
      </div>
    </section>
  </main>

  <footer class="site-footer">
    <div class="container footer-grid">
      <div>
        <h3>Brojogopal Sapui</h3>
        <p>AI Security• Hardware Trust • Edge/Physical AI</p>
      </div>
      <div>
        <h4>Main Pages</h4>
        <a href="../research.html">Research</a>
        <a href="../ongoing-work.html">Trending Topics</a>
        <a href="../publications.html">Resources</a>
      </div>
      <div>
        <h4>Focus</h4>
        <p>Cross-layer AI security, trustworthy deployment, hardware-aware defense, and physical intelligence.</p>
      </div>
    </div>
  </footer>

  <script src="../assets/js/main.js"></script>
</body>
</html>'''.strip()


def update_ongoing_work_html(input_html_path: Path, article_id: str, article_html: str, replace_duplicate: bool = True) -> Path:
    html_text = input_html_path.read_text(encoding='utf-8')
    doctype = detect_doctype(html_text)
    soup = BeautifulSoup(html_text, 'html.parser')

    watch_stack = soup.select_one('div.watch-stack')
    if watch_stack is None:
        raise ValueError('Could not find div.watch-stack in ongoing-work.html')

    if replace_duplicate:
        existing = watch_stack.find('article', {'id': article_id})
        if existing:
            existing.decompose()

    fragment = BeautifulSoup(article_html, 'html.parser')
    new_article = fragment.find('article')
    first_existing = watch_stack.find('article')
    if first_existing:
        first_existing.insert_before('\n')
        first_existing.insert_before(new_article)
        first_existing.insert_before('\n\n          ')
    else:
        watch_stack.append(new_article)

    input_html_path.write_text(doctype + '\n' + str(soup), encoding='utf-8')
    return input_html_path


def update_index_html_current_structure(input_html_path: Path, post_id: str, title: str, preview: str) -> Path:
    html_text = input_html_path.read_text(encoding='utf-8')
    doctype = detect_doctype(html_text)
    soup = BeautifulSoup(html_text, 'html.parser')

    alert = soup.select_one('div.research-alert')
    if alert is None:
        raise ValueError(
            'Could not find div.research-alert in index.html. This updater is intentionally limited to the current homepage structure.'
        )

    chips = alert.select_one('div.hero-chips')
    if chips is None:
        raise ValueError('Could not find div.hero-chips inside div.research-alert')

    chips.clear()
    span = soup.new_tag('span')
    span.string = title if _clean(title) else preview
    chips.append(span)

    link = alert.select_one('a.research-alert-link')
    if link is None:
        raise ValueError('Could not find a.research-alert-link inside div.research-alert')
    link['href'] = f'ongoing-work.html#{post_id}'
    link.string = 'Read the research watch →'

    input_html_path.write_text(doctype + '\n' + str(soup), encoding='utf-8')
    return input_html_path


def write_full_post_html(output_path: Path, full_post_html: str) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(full_post_html, encoding='utf-8')
    return output_path


def update_site(
    repo_root: Path,
    docx_path: Path | None = None,
    auto_pick_latest: bool = True,
    weekly_inputs_dir: str = 'weekly-inputs',
    posts_dir: str = 'posts',
    ongoing_html: str = 'ongoing-work.html',
    index_html: str = 'index.html',
    replace_duplicate: bool = True,
) -> UpdateResult:
    repo_root = Path(repo_root).resolve()
    weekly_dir = repo_root / weekly_inputs_dir
    posts_path = repo_root / posts_dir
    ongoing_path = repo_root / ongoing_html
    index_path = repo_root / index_html
    template_hints = {
        'weekly-research-watch-template.docx',
        'template.docx',
        'research-watch-template.docx',
    }

    if docx_path is None:
        if auto_pick_latest:
            chosen_docx = pick_latest_docx(weekly_dir, template_hints)
        else:
            raise ValueError('docx_path must be provided when auto_pick_latest=False')
    else:
        chosen_docx = Path(docx_path)
        if not chosen_docx.is_absolute():
            chosen_docx = (repo_root / chosen_docx).resolve()

    if not chosen_docx.exists():
        raise FileNotFoundError(f'DOCX file not found: {chosen_docx}')
    if not ongoing_path.exists():
        raise FileNotFoundError(f'File not found: {ongoing_path}')
    if not index_path.exists():
        raise FileNotFoundError(f'File not found: {index_path}')

    data = parse_weekly_docx(chosen_docx)
    _, normalized_stem = normalize_docx_data(data, chosen_docx)
    post_id, title, preview, article_html = build_watch_article(data, normalized_stem)

    full_post_html = build_full_post_html(data, normalized_stem)
    post_output_path = posts_path / f'{normalized_stem}.html'
    write_full_post_html(post_output_path, full_post_html)
    update_ongoing_work_html(ongoing_path, post_id, article_html, replace_duplicate=replace_duplicate)
    update_index_html_current_structure(index_path, post_id, title, preview)

    return UpdateResult(
        docx_path=chosen_docx,
        post_id=post_id,
        title=title,
        preview=preview,
        ongoing_path=ongoing_path,
        index_path=index_path,
        post_path=post_output_path,
    )
